from docx import Document
import os

def convert_text_to_word(text: str, filename: str) -> str:
    doc = Document()
    doc.add_heading("PHDCIScanner - Converted Document", 0)
    
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    output_path = f"outputs/{filename}.docx"
    doc.save(output_path)
    return output_path