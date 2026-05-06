from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas
from docx import Document
from docx.shared import Inches
import openpyxl
import io
import os

os.makedirs("outputs", exist_ok=True)

def image_to_pdf(image_bytes: bytes, filename: str) -> str:
    output_path = f"outputs/{filename}.pdf"
    image = Image.open(io.BytesIO(image_bytes))
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    # Resize image to fit A4
    a4_width, a4_height = A4  # 595 x 842 points
    max_w = a4_width - 40
    max_h = a4_height - 40
    
    ratio = min(max_w / image.width, max_h / image.height)
    new_w = int(image.width * ratio)
    new_h = int(image.height * ratio)
    image = image.resize((new_w, new_h), Image.LANCZOS)
    
    temp_path = f"outputs/{filename}_temp.jpg"
    image.save(temp_path, "JPEG")
    
    # Draw directly on canvas
    c = pdf_canvas.Canvas(output_path, pagesize=A4)
    x = (a4_width - new_w) / 2
    y = (a4_height - new_h) / 2
    c.drawImage(temp_path, x, y, width=new_w, height=new_h)
    c.save()
    return output_path

def image_to_word(image_bytes: bytes, filename: str) -> str:
    output_path = f"outputs/{filename}.docx"
    image = Image.open(io.BytesIO(image_bytes))
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    max_px = 1000
    if image.width > max_px:
        ratio = max_px / image.width
        image = image.resize((int(image.width * ratio), int(image.height * ratio)), Image.LANCZOS)
    
    temp_path = f"outputs/{filename}_temp.jpg"
    image.save(temp_path, "JPEG")
    
    doc = Document()
    doc.add_heading('PHDCIScanner - Converted Image', 0)
    max_width = 6.0
    aspect = image.height / image.width
    new_height = max_width * aspect
    if new_height > 8.0:
        new_height = 8.0
        max_width = new_height / aspect
    doc.add_picture(temp_path, width=Inches(max_width))
    doc.save(output_path)
    return output_path

def image_to_excel(image_bytes: bytes, filename: str) -> str:
    output_path = f"outputs/{filename}.xlsx"
    image = Image.open(io.BytesIO(image_bytes))
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    max_px = 800
    if image.width > max_px:
        ratio = max_px / image.width
        image = image.resize((int(image.width * ratio), int(image.height * ratio)), Image.LANCZOS)
    
    temp_path = f"outputs/{filename}_temp.jpg"
    image.save(temp_path, "JPEG")
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Image"
    ws['A1'] = 'PHDCIScanner - Converted Image'
    img = openpyxl.drawing.image.Image(temp_path)
    img.anchor = 'A3'
    ws.add_image(img)
    wb.save(output_path)
    return output_path